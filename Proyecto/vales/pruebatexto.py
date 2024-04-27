import os
import win32com.client
import shutil
from docx import Document
from docx.shared import Pt
from tools.tools import Error, detectar_horario

def borrar_archivo(ruta):
    try:
        os.remove(ruta)
    except FileNotFoundError:
        a = f"El archivo {ruta} no existe."
        Error(a)
    except Exception as e:
        a = f"No se pudo borrar el archivo {ruta}: {e}"
        Error(a)

def reemplazar_palabra_en_documento(archivo_docx, C1, C2, C3, C4, C5, C6, Herra, Medi, Canti, observaciones, C1N, C2N, C3N, C4N, C5N, C6N, HerraN, MediN, CantiN, MotivoN, archivo_salida):
    # Abrir el documento existente
    doc = Document(archivo_docx)
    abc = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M", "N", "O", "P", "Q", "R", "S", "T"]
    X = 0
    Y = 1
    # Recorrer cada párrafo en el documento
    for paragraph in doc.paragraphs:
        # Buscar y reemplazar la palabra en el texto de cada párrafo
        if C1 in paragraph.text:
            paragraph.text = paragraph.text.replace(C1, C1N)
            # Aplicar estilo al texto reemplazado
            run = paragraph.runs[-1]  # Obtener el último Run (fragmento de texto) en el párrafo
            run.font.size = Pt(18)  # Establecer el tamaño de fuente en 18 puntos
            run.font.nama = "Aptos Display"

        if C2 in paragraph.text:
            paragraph.text = paragraph.text.replace(str(C2), str(C2N))
            run = paragraph.runs[-1]
            run.font.size = Pt(16)
            run.font.name = "Bahnschrift"
        if C3 in paragraph.text:
            paragraph.text = paragraph.text.replace(str(C3), str(C3N))
            run = paragraph.runs[-1]
            run.font.size = Pt(16)
            run.font.name = "Bahnschrift"
        if C6 in paragraph.text:
            paragraph.text = paragraph.text.replace(C6, str(C6N).zfill(5))
            run = paragraph.runs[-1]
            run.font.size = Pt(16)
            run.font.name = "Bahnschrift"
        if C4 in paragraph.text:
            paragraph.text = paragraph.text.replace(C4, C4N)
            run = paragraph.runs[-1]
            run.font.size = Pt(14)
            run.font.name = "Bahnschrift"
        if C5 in paragraph.text:
            paragraph.text = paragraph.text.replace(C5, C5N)
            run = paragraph.runs[-1]
            run.font.size = Pt(14)
            run.font.name = "Bahnschrift"
        for table in doc.tables:
            # Iterar sobre las filas en la tabla
            for row in table.rows:
                # Iterar sobre las celdas en la fila
                for cell in row.cells:
                    # Iterar sobre los párrafos en la celda
                    for paragraph in cell.paragraphs:
                        if Y == 1:
                            # Verificar si el texto está en el párrafo
                            if (Herra + abc[X]) in paragraph.text:
                                # Reemplazar
                                paragraph.text = paragraph.text.replace(Herra + abc[X], HerraN[X])
                                # Modificar el estilo de la última ejecución en el párrafo
                                run = paragraph.runs[-1]
                                run.font.size = Pt(11)
                                run.font.name = "Calibri (Cuerpo)"
                            elif (Medi + abc[X]) in paragraph.text:
                                paragraph.text = paragraph.text.replace(Medi + abc[X], MediN[X])
                                run = paragraph.runs[-1]
                                run.font.size = Pt(11)
                                run.font.name = "Calibri (Cuerpo)"
                            elif (Canti + abc[X]) in paragraph.text:
                                paragraph.text = paragraph.text.replace(Canti + abc[X], CantiN[X])
                                run = paragraph.runs[-1]
                                run.font.size = Pt(11)
                                run.font.name = "Calibri (Cuerpo)"
                            elif (observaciones + abc[X]) in paragraph.text:
                                paragraph.text = paragraph.text.replace(observaciones + abc[X], MotivoN)
                                run = paragraph.runs[-1]
                                run.font.size = Pt(11)
                                run.font.name = "Calibri (Cuerpo)"
                                X += 1

                            if X == len(HerraN) and X == len(MediN) and X == len(CantiN):
                                Y = 2

                        elif Y == 2:
                            if (Herra + abc[X]) in paragraph.text:
                                # Reemplazar
                                paragraph.text = paragraph.text.replace(Herra + abc[X], "---")
                                # Modificar el estilo de la última ejecución en el párrafo
                                run = paragraph.runs[-1]
                                run.font.size = Pt(11)
                                run.font.name = "Calibri (Cuerpo)"
                            elif (Medi + abc[X]) in paragraph.text:
                                paragraph.text = paragraph.text.replace(Medi + abc[X], "---")
                                run = paragraph.runs[-1]
                                run.font.size = Pt(11)
                                run.font.name = "Calibri (Cuerpo)"
                            elif (Canti + abc[X]) in paragraph.text:
                                paragraph.text = paragraph.text.replace(Canti + abc[X], "---")
                                run = paragraph.runs[-1]
                                run.font.size = Pt(11)
                                run.font.name = "Calibri (Cuerpo)"
                            elif (observaciones + abc[X]) in paragraph.text:
                                paragraph.text = paragraph.text.replace(observaciones + abc[X], "---")
                                run = paragraph.runs[-1]
                                run.font.size = Pt(11)
                                run.font.name = "Calibri (Cuerpo)"
                                X += 1

                            if X == len(abc):
                                Y = 0

    # Guardar los cambios en el mismo archivo
    doc.save(archivo_salida)

def vales(HerraN, MediN, cantiN, C2N, C5N, C4N, C6N, MotivoN):
    # Determinar el término correcto para "Escritorio" según el sistema operativo
    escritorio = os.path.expanduser("~\\Desktop") if os.name == "nt" else os.path.expanduser("~\\Escritorio")

    # Ruta al directorio del escritorio
    desktop_path = os.path.join(escritorio, "Historial de vales")

    # Nombre del archivo .docx existente
    archivo_docx = os.path.join(desktop_path, "Prueba.docx")
    print(archivo_docx)
    # Palabra que quieres reemplazar en el documento
    C1 = "cambio1"
    C2 = "cambio2"
    C3 = "cambio3"
    C4 = "cambio4"
    C5 = "cambio5"
    C6 = "cambio6"
    Herra = "Nombre "
    Medi = "Medida "
    Canti = "E"
    observaciones="Observaciones "
    # Nueva palabra con la que deseas reemplazar
    C1N = "CONSUMIBLES"
    C3N = detectar_horario()

    # Nombre del nuevo archivo .docx de salida
    archivo_salida = os.path.join(desktop_path, f"Vale-{C6N}.docx")

    # Reemplazar la palabra en el documento existente y guardar los cambios en un nuevo archivo
    reemplazar_palabra_en_documento(archivo_docx, C1, C2, C3, C4, C5, C6, Herra, Medi, Canti,observaciones, C1N, C2N, C3N, C4N, C5N, C6N, HerraN, MediN, cantiN, MotivoN, archivo_salida)

    # Convertir el archivo .docx a .pdf
    wdFormatPDF = 17
    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(os.path.abspath(archivo_salida))
    doc.SaveAs(os.path.join(desktop_path, f"Vale-{C6N}.pdf"), FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

    # Borrar el archivo .docx
    borrar_archivo(os.path.abspath(archivo_salida))
